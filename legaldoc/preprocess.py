from pandas import *
from glob import glob
from docx import Document
from pathlib import Path


def set_data_path(base_path:str, pattern_str:str)->glob:
    """
    Make glob obj from string
    :param base_path: project path
    :param pattern_str: from project path to data location and append pattern(s)
    :return: glob object
    """
    return glob(base_path+ '/' + pattern_str, recursive=True)


def load_docx(doc_path:glob)->list:
    """
    Read and convert word doc to data dict
       Word supports a section, a division of a document having the same page layout settings,
       such as margins and page orientation.
    This is how, a document can contain some pages in portrait and others in landscape.

    :param doc_path: file location
    :return: list of doc data
    """
    docs = []
    for ndx, file in enumerate(doc_path):
        document = Document(file)

        docs.append({
            'ndx': ndx,
            'ndx_name': file.split('/')[-1],
            'doc_type': file.split('.')[1],
            'raw': [ raw.text.replace('\xa0', ' ') for raw in document.paragraphs ],
            'tables': document.tables,
        })
    return docs


def clean_up(docs):
    for doc in docs:
        doc['processed'] = [ raw.replace('\xa0', ' ') for raw in doc['raw'] ]
    pass


def review_tables(doc_tables: list)->list:
    """
    Print out contents of tables in text
    :param doc_tables: list of Table objects
    :return: list of table contents
    """
    res = []
    for i, tbl in enumerate(doc_tables):
        for r in range(len(tbl.rows)):
            for c in range(len(tbl.columns)):
                res.append((i, r, c, tbl.cell(r, c).text.replace('\xa0', ' ')))
    return res

